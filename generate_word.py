from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

PRIMARY = RGBColor(0x1a, 0x1a, 0x2e)
HIGHLIGHT = RGBColor(0x0f, 0x34, 0x60)
GOLD = RGBColor(0xe9, 0x45, 0x60)
SUBTLE = RGBColor(0x6c, 0x75, 0x7d)
DARK = RGBColor(0x2d, 0x2d, 0x2d)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_horizontal_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="dee2e6"/></w:pBdr>')
    pPr.append(pBdr)

def build_word():
    doc = Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # Title
    p = doc.add_paragraph()
    run = p.add_run("HARSH")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = PRIMARY
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    run = p.add_run("Senior Mobile & Full Stack Engineer")
    run.font.size = Pt(13)
    run.font.color.rgb = HIGHLIGHT
    p.paragraph_format.space_after = Pt(8)

    # Summary
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    summary = (
        "Results-driven Senior Engineer with 8.5+ years of hands-on experience architecting and delivering "
        "enterprise-grade mobile and full-stack applications. Specialized in cross-platform mobile development using "
        "Xamarin.Forms, .NET MAUI, and Flutter, coupled with robust backend systems built on ASP.NET Core, "
        "FastAPI, SQL Server, and Supabase. Passionate about integrating AI/ML capabilities into production "
        "applications using Python, Ollama, and modern LLM frameworks. Proven track record across healthcare, "
        "hospitality, logistics, finance, and real-time communication domains."
    )
    run = p.add_run(summary)
    run.font.size = Pt(10)
    run.font.color.rgb = DARK

    # Stats table
    stats_table = doc.add_table(rows=2, cols=4)
    stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    stats = [("8.5+", "Years Experience"), ("10+", "Projects Delivered"), ("5", "Platforms"), ("3+", "AI Projects")]
    for i, (val, label) in enumerate(stats):
        cell = stats_table.rows[0].cells[i]
        set_cell_shading(cell, "f8f9fa")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val)
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = PRIMARY

        cell = stats_table.rows[1].cells[i]
        set_cell_shading(cell, "f8f9fa")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.font.size = Pt(8)
        run.font.color.rgb = SUBTLE

    doc.add_paragraph()
    add_horizontal_line(doc)

    # Career Trajectory
    p = doc.add_paragraph()
    run = p.add_run("CAREER TRAJECTORY")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = PRIMARY
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)

    experiences = [
        ("Senior Software Engineer", "Jan 2026 – Present", "Leading architecture & development of enterprise mobile apps, AI integration, and full-stack solutions."),
        ("Xamarin Developer", "Jan 2022 – Dec 2025", "Designed and delivered cross-platform mobile applications for healthcare, logistics, and hospitality using Xamarin.Forms & .NET MAUI."),
        ("Associate Software Engineer", "Jun 2017 – Dec 2021", "Built foundational mobile and backend systems, REST APIs, database design, and client-facing applications."),
    ]

    for role, period, desc in experiences:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(f"▸ {role}")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = DARK

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(period)
        run.font.size = Pt(9)
        run.font.color.rgb = SUBTLE

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(desc)
        run.font.size = Pt(9.5)
        run.font.color.rgb = DARK

    add_horizontal_line(doc)

    # Skills
    p = doc.add_paragraph()
    run = p.add_run("CORE SKILLS & TECHNOLOGIES")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = PRIMARY
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)

    skills = [
        ("MOBILE", ".NET MAUI · Xamarin.Forms · Flutter · Dart · C# · MVVM"),
        ("BACKEND", "ASP.NET Core · FastAPI · REST APIs · Entity Framework · Python"),
        ("AI / ML", "Ollama · LLM Integration · AI Pipelines · Computer Vision"),
        ("CLOUD & DB", "SQL Server · Supabase · SQLite · Azure · Serverless"),
        ("TOOLS", "Visual Studio · VS Code · Git · Postman · Agile / Scrum"),
        ("OTHER", "HIPAA · HL7 · Real-time Chat · BLE · Cross-platform"),
    ]

    skills_table = doc.add_table(rows=len(skills), cols=2)
    for i, (cat, techs) in enumerate(skills):
        cell = skills_table.rows[i].cells[0]
        p = cell.paragraphs[0]
        run = p.add_run(cat)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = HIGHLIGHT

        cell = skills_table.rows[i].cells[1]
        p = cell.paragraphs[0]
        run = p.add_run(techs)
        run.font.size = Pt(9)
        run.font.color.rgb = DARK

    # Skill bars as text representation
    doc.add_paragraph()
    skill_bars = [
        ("Xamarin / .NET MAUI", 95),
        ("Flutter / Dart", 88),
        ("ASP.NET Core / C#", 92),
        ("Python / FastAPI / AI", 82),
        ("SQL Server / Supabase", 88),
        ("UI/UX & Cross-Platform", 85),
    ]
    for label, pct in skill_bars:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        bar_filled = "█" * (pct // 5)
        bar_empty = "░" * (20 - pct // 5)
        run = p.add_run(f"{label:<25}")
        run.font.size = Pt(9)
        run.font.color.rgb = DARK
        run = p.add_run(f" {bar_filled}{bar_empty} ")
        run.font.size = Pt(9)
        run.font.color.rgb = HIGHLIGHT
        run = p.add_run(f"{pct}%")
        run.font.size = Pt(8)
        run.font.color.rgb = SUBTLE

    # Page break for projects
    doc.add_page_break()

    # Project Portfolio
    p = doc.add_paragraph()
    run = p.add_run("PROJECT PORTFOLIO (10 Projects)")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = PRIMARY
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    run = p.add_run(
        "A curated collection of production applications spanning healthcare, AI, hospitality, logistics, and finance — "
        "demonstrating depth in mobile engineering, backend architecture, and AI integration."
    )
    run.font.size = Pt(8)
    run.font.color.rgb = SUBTLE
    p.paragraph_format.space_after = Pt(8)

    projects = [
        {
            "num": 1, "title": "HealthxApp for Patients",
            "subtitle": "Enterprise Healthcare Platform – Patient Portal",
            "desc": "A comprehensive patient-facing mobile application enabling appointment scheduling, secure document management, HIPAA-compliant video consultations, and real-time chat with healthcare providers.",
            "bullets": [
                "Built end-to-end appointment booking with provider availability and calendar sync",
                "Implemented secure video consultation using WebRTC with encrypted streams",
                "Developed real-time chat system with message history and push notifications",
                "Integrated HL7/FHIR standards for healthcare data interoperability",
                "Document upload/download with secure cloud storage and access control",
                "HIPAA-compliant data handling with encryption at rest and in transit",
                "Offline-first architecture with background sync for poor connectivity scenarios",
                "Role-based access control separating patient, provider, and admin flows",
            ],
            "tech": "Xamarin.Forms · C# · ASP.NET Core API · SQL Server · HL7 · HIPAA · MVVM · Push Notifications",
        },
        {
            "num": 2, "title": "HealthxApp for Providers",
            "subtitle": "Enterprise Healthcare Platform – Provider Dashboard",
            "desc": "Provider-facing companion application enabling doctors and clinical staff to manage appointments, review patient documents, conduct video consultations, and communicate securely with patients.",
            "bullets": [
                "Appointment management dashboard with schedule overview and conflict detection",
                "Patient document review system with annotation and approval workflows",
                "Video consultation initiation with recording capability for medical records",
                "Secure messaging with read receipts, typing indicators, and file sharing",
                "Integration with hospital EMR systems via HL7 interfaces",
                "Multi-location support for providers practising across clinics",
                "Analytics dashboard showing consultation metrics and patient engagement",
                "Notification system for urgent patient requests and schedule changes",
            ],
            "tech": "Xamarin.Forms · C# · ASP.NET Core API · SQL Server · HL7 · HIPAA · MVVM · WebRTC",
        },
        {
            "num": 3, "title": "Clipboard",
            "subtitle": "Patient Intake & Clinical Workflow Management",
            "desc": "Digital patient intake solution replacing paper-based processes — handling appointments, consent forms, questionnaires, and medical history management with seamless offline support.",
            "bullets": [
                "Dynamic consent form builder with e-signature capture and PDF generation",
                "Configurable medical questionnaires with conditional logic and branching",
                "Complete medical history intake with allergies, medications, and conditions",
                "Offline-first SQLite storage with intelligent background synchronisation",
                "Migrated from Xamarin.Forms to .NET MAUI for long-term platform support",
                "Barcode scanning for patient identification and record lookup",
                "Multi-language support for diverse patient demographics",
                "Seamless integration with backend .NET Core APIs and SQL Server",
            ],
            "tech": "Xamarin.Forms · .NET MAUI · C# · SQLite · SQL Server · .NET Core API · MVVM · Offline Sync",
        },
        {
            "num": 4, "title": "SMD Mobile",
            "subtitle": "Hospital & Clinic Management System",
            "desc": "Comprehensive hospital and clinic management mobile application covering patient records, medicine inventory, appointment scheduling, and clinical workflow automation.",
            "bullets": [
                "Patient registration and record management with search and filtering",
                "Medicine inventory tracking with stock alerts and expiry management",
                "Appointment scheduling with provider calendars and automated reminders",
                "Migrated entire codebase from Xamarin.Forms to .NET MAUI",
                "Improved startup time and memory efficiency via framework optimisation",
                "REST API integration for secure, real-time backend communication",
                "Resolved Android/iOS platform-specific UI and compatibility issues",
                "Extensive testing, debugging, and production performance tuning",
            ],
            "tech": "Xamarin.Forms · .NET MAUI · C# · SQLite · SQL Server · .NET Core API · MVVM · REST APIs",
        },
        {
            "num": 5, "title": "LogicWash",
            "subtitle": "Commercial Car Wash Management Solution",
            "desc": "A streamlined mobile solution for managing commercial car wash operations — handling service queues, customer management, wash tracking, and operational analytics.",
            "bullets": [
                "Service queue management with real-time status tracking",
                "Customer profile management with wash history and preferences",
                "Multiple wash package configuration with pricing tiers",
                "Offline-capable SQLite storage for uninterrupted operations",
                "Staff scheduling and performance tracking dashboards",
                "Revenue reporting with daily, weekly, and monthly breakdowns",
                "Push notifications for service completion and promotional offers",
                "Clean MVVM architecture ensuring maintainability and testability",
            ],
            "tech": ".NET MAUI · C# · SQLite · MVVM · Local Storage · Cross-platform",
        },
        {
            "num": 6, "title": "Digital Butler",
            "subtitle": "AI-Powered Hotel & Resort Guest Management Platform",
            "desc": "An intelligent hospitality platform leveraging AI to provide personalised guest experiences — from smart room booking and concierge services to AI-driven recommendations and guest communication.",
            "bullets": [
                "AI-powered concierge chatbot using Ollama LLM for natural guest interactions",
                "Smart room booking with dynamic pricing and availability optimisation",
                "Guest preference learning system for personalised service recommendations",
                "Real-time communication between guests and hotel staff via chat",
                "Integration with hotel PMS for room status, housekeeping, and maintenance",
                "Multi-property support with centralised management dashboard",
                "Supabase real-time database for instant updates across all clients",
                "Python FastAPI backend with AI pipeline orchestration",
            ],
            "tech": "Flutter · Dart · Supabase · Python · FastAPI · AI · Ollama · Real-time DB · Cross-platform",
        },
        {
            "num": 7, "title": "SmartLogix",
            "subtitle": "AI-Powered Shipment & Logistics Tracking Platform",
            "desc": "Intelligent logistics platform that uses AI for shipment registration, route optimisation, delivery prediction, and real-time tracking — transforming traditional shipping workflows.",
            "bullets": [
                "AI-driven shipment registration with automatic data extraction from documents",
                "Real-time GPS tracking with estimated delivery time predictions",
                "Route optimisation engine using ML models for cost and time efficiency",
                "Automated status updates with intelligent notification triggers",
                "Document scanning and OCR for bills of lading and shipping labels",
                "Analytics dashboard with shipment performance metrics and trends",
                "Multi-carrier integration with unified tracking interface",
                "Supabase real-time sync for instant status propagation",
            ],
            "tech": "Flutter · Dart · Supabase · Python · FastAPI · AI · Ollama · OCR · Real-time Tracking",
        },
        {
            "num": 8, "title": "ReconVision",
            "subtitle": "AI Platform for Crime & Anomaly Detection",
            "desc": "Advanced AI-powered surveillance and analytics platform combining computer vision for crime detection, facial recognition, violence detection, and behavioural analytics in real-time video feeds.",
            "bullets": [
                "Real-time video feed analysis with computer vision models for anomaly detection",
                "Face detection and recognition with configurable alert thresholds",
                "Violence and aggressive behaviour detection using trained ML models",
                "Analytics dashboard with heat maps, incident timelines, and trend analysis",
                "Multi-camera support with centralised monitoring interface",
                "Alert system with configurable escalation workflows",
                "Historical incident search with video playback and evidence export",
                "Python AI pipeline with Ollama integration for contextual analysis",
            ],
            "tech": "Flutter · Dart · Supabase · Python · FastAPI · AI · Ollama · Computer Vision · Analytics",
        },
        {
            "num": 9, "title": "Pinglet",
            "subtitle": "Real-Time Chat Application",
            "desc": "A modern, performant real-time messaging application with instant message delivery, presence indicators, media sharing, and group conversation support.",
            "bullets": [
                "Real-time messaging with sub-second delivery using Supabase Realtime",
                "Online/offline presence indicators with last-seen timestamps",
                "Media sharing including images, documents, and voice messages",
                "Group chat creation and management with admin controls",
                "Message read receipts and typing indicators",
                "Push notifications for background message delivery",
                "End-to-end message history with search functionality",
                "Clean Material Design UI with smooth animations",
            ],
            "tech": "Flutter · Dart · Supabase · Realtime · Push Notifications · Cross-platform",
        },
        {
            "num": 10, "title": "Finance App",
            "subtitle": "Personal Finance & Expense Tracking",
            "desc": "A personal finance management application enabling users to track expenses, set budgets, visualise spending patterns, and gain insights into their financial habits.",
            "bullets": [
                "Expense recording with category tagging and receipt photo capture",
                "Budget creation with configurable limits and overspend alerts",
                "Interactive charts and graphs for spending pattern visualisation",
                "Monthly and yearly financial summaries with export capability",
                "Recurring expense tracking with automatic categorisation",
                "Multi-currency support with real-time conversion rates",
                "Secure authentication with biometric login support",
                "Cloud sync via Supabase for multi-device access",
            ],
            "tech": "Flutter · Dart · Supabase · Charts · Cross-platform · Biometric Auth",
        },
    ]

    for proj in projects:
        # Title
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(f"{proj['num']}  ")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = GOLD
        run = p.add_run(proj["title"])
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = PRIMARY

        # Subtitle
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.7)
        run = p.add_run(proj["subtitle"])
        run.font.size = Pt(9)
        run.font.color.rgb = SUBTLE

        # Description
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(proj["desc"])
        run.font.size = Pt(9.5)
        run.font.color.rgb = DARK

        # Bullets in 2-column table
        half = len(proj["bullets"]) // 2
        bullet_table = doc.add_table(rows=half, cols=2)
        for i in range(half):
            cell_l = bullet_table.rows[i].cells[0]
            p = cell_l.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run("● ")
            run.font.size = Pt(8)
            run.font.color.rgb = GOLD
            run = p.add_run(proj["bullets"][i])
            run.font.size = Pt(8.5)
            run.font.color.rgb = DARK

            cell_r = bullet_table.rows[i].cells[1]
            p = cell_r.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run("● ")
            run.font.size = Pt(8)
            run.font.color.rgb = GOLD
            run = p.add_run(proj["bullets"][half + i])
            run.font.size = Pt(8.5)
            run.font.color.rgb = DARK

        # Tech line
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run("Tech: ")
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = HIGHLIGHT
        tech_tags = " · ".join([f"[{t.strip()}]" for t in proj["tech"].split("·")])
        run = p.add_run(tech_tags)
        run.font.size = Pt(8)
        run.font.color.rgb = HIGHLIGHT

        add_horizontal_line(doc)

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "This portfolio represents a curated selection of production applications. "
        "Each project demonstrates expertise in architecture design, clean code practices, "
        "and delivering real business value through technology."
    )
    run.font.size = Pt(8)
    run.font.color.rgb = SUBTLE

    output = "/Users/harshshah92/Documents/HarshShah_Portfolio/Harsh_Portfolio.docx"
    doc.save(output)
    print(f"✅ Word document generated: {output}")

if __name__ == "__main__":
    build_word()
